#!/usr/bin/env python3
"""
Benchmark script to build and test all templates.
Generates a comparison report of original vs. optimized templates.
"""

import json
import subprocess
import sys
from pathlib import Path
from datetime import datetime

# Try to import python-docx for Word document generation
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Templates to benchmark
ORIGINAL_TEMPLATES = ["avi", "gif", "bmp", "jpg", "midi", "mp3", "mp4", "png"]
OPTIMIZED_TEMPLATES = ["avi-orig", "gif-orig", "bmp-orig", "jpg-orig",
                       "midi-orig", "mp3-orig", "mp4-orig", "png-orig"]

# Fuzzing parameters
FUZZ_COUNT = 10000
WORKDIR = Path(".")

def build_fuzzer(template_name):
    """Build fuzzer for a given template."""
    print(f"\n{'='*60}")
    print(f"Building fuzzer for: {template_name}")
    print(f"{'='*60}")

    try:
        result = subprocess.run(
            ["bash", "build_new.sh", template_name],
            cwd=WORKDIR,
            capture_output=True,
            timeout=120,
            text=True
        )

        if result.returncode != 0:
            print(f"❌ Build failed for {template_name}")
            print(f"stderr: {result.stderr[-500:]}")  # Last 500 chars
            return False

        print(f"✅ Built successfully: {template_name}-fuzzer")
        return True

    except subprocess.TimeoutExpired:
        print(f"❌ Build timeout for {template_name}")
        return False
    except Exception as e:
        print(f"❌ Build error for {template_name}: {e}")
        return False


def run_fuzz_manager(filetype):
    """Run fuzz_manager for a given filetype."""
    print(f"\n  Fuzzing {filetype} ({FUZZ_COUNT} files)...")

    try:
        result = subprocess.run(
            ["python3", "scripts/fuzz_manager.py", filetype, str(FUZZ_COUNT)],
            cwd=WORKDIR,
            capture_output=True,
            timeout=600,  # 10 minutes
            text=True
        )

        if result.returncode != 0:
            print(f"  ⚠️  Fuzzing had issues")
            print(f"  Last output: {result.stdout[-300:]}")
            return None

        # Extract timestamp from output
        for line in result.stdout.split('\n'):
            if 'Timestamp' in line:
                timestamp = line.split(':')[1].strip()
                return timestamp

        return None

    except subprocess.TimeoutExpired:
        print(f"  ❌ Fuzzing timeout for {filetype}")
        return None
    except Exception as e:
        print(f"  ❌ Fuzzing error for {filetype}: {e}")
        return None


def read_report(filetype, timestamp):
    """Read the generated report.json file."""
    report_path = WORKDIR / "output" / filetype / timestamp / "report.json"

    try:
        with open(report_path, 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"  ⚠️  Report not found: {report_path}")
        return None
    except Exception as e:
        print(f"  ⚠️  Error reading report: {e}")
        return None


def create_word_report(results, filename="BENCHMARK_REPORT.docx"):
    """Create Word document report from results."""
    if not DOCX_AVAILABLE:
        print(f"[warn] python-docx not installed; skipping Word document generation")
        return False

    doc = Document()

    # Title
    title = doc.add_heading('Fuzzer Benchmarking Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Generated date
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()

    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('This report compares the performance of original templates vs. optimized templates.')
    doc.add_paragraph()

    # Statistics table
    doc.add_heading('Statistics', level=2)
    table = doc.add_table(rows=1, cols=11)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Template', 'Type', 'Count', 'Generated', 'Valid', 'Invalid',
               'Gen Speed (f/s)', 'Gen Time (s)', 'Val Time (s)', 'Total (s)', 'Valid %']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True

    # Data rows
    for entry in sorted(results, key=lambda x: x['template']):
        row_cells = table.add_row().cells
        row_cells[0].text = entry['template']
        row_cells[1].text = entry['type']
        row_cells[2].text = str(entry['count'])
        row_cells[3].text = str(entry['generated'])
        row_cells[4].text = str(entry['valid'])
        row_cells[5].text = str(entry['invalid'])
        row_cells[6].text = f"{entry['gen_speed']:.2f}"
        row_cells[7].text = f"{entry['gen_time']:.2f}"
        row_cells[8].text = f"{entry['val_time']:.2f}"
        row_cells[9].text = f"{entry['total_time']:.2f}"
        row_cells[10].text = f"{entry['valid_rate']:.2f}%"

    doc.add_paragraph()

    # Comparison section
    doc.add_heading('Comparison', level=1)

    original = [r for r in results if r['type'] == 'Original']
    optimized = [r for r in results if r['type'] == 'Optimized']

    if original and optimized:
        avg_gen_orig = sum(r['gen_speed'] for r in original) / len(original)
        avg_gen_opt = sum(r['gen_speed'] for r in optimized) / len(optimized)
        avg_val_orig = sum(r['val_time'] for r in original) / len(original)
        avg_val_opt = sum(r['val_time'] for r in optimized) / len(optimized)
        avg_valid_orig = sum(r['valid_rate'] for r in original) / len(original)
        avg_valid_opt = sum(r['valid_rate'] for r in optimized) / len(optimized)

        doc.add_heading('Generation Performance', level=2)
        doc.add_paragraph(f'Original templates avg speed: {avg_gen_orig:.2f} files/sec')
        doc.add_paragraph(f'Optimized templates avg speed: {avg_gen_opt:.2f} files/sec')
        speedup = avg_gen_opt / avg_gen_orig if avg_gen_orig > 0 else 0
        doc.add_paragraph(f'Speedup: {speedup:.2f}x')
        doc.add_paragraph()

        doc.add_heading('Validation Performance', level=2)
        doc.add_paragraph(f'Original templates avg time: {avg_val_orig:.2f} seconds')
        doc.add_paragraph(f'Optimized templates avg time: {avg_val_opt:.2f} seconds')
        reduction = ((avg_val_orig - avg_val_opt) / avg_val_orig * 100) if avg_val_orig > 0 else 0
        doc.add_paragraph(f'Reduction: {reduction:.1f}%')
        doc.add_paragraph()

        doc.add_heading('Validity Rates', level=2)
        doc.add_paragraph(f'Original templates avg valid rate: {avg_valid_orig:.2f}%')
        doc.add_paragraph(f'Optimized templates avg valid rate: {avg_valid_opt:.2f}%')

    # Save document
    doc.save(filename)
    return True


def format_report_for_markdown(results):
    """Format results as a markdown table."""
    markdown = """# Fuzzer Benchmarking Report

Generated: {timestamp}

## Summary

This report compares the performance of original templates vs. optimized templates.

### Statistics

""".format(timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    # Table header
    markdown += """
| Template | Type | Count | Generated | Valid | Invalid | Gen Speed (f/s) | Gen Time (s) | Val Time (s) | Total Time (s) |
|----------|------|-------|-----------|-------|---------|-----------------|--------------|--------------|----------------|
"""

    for entry in results:
        row = f"""| {entry['template']} | {entry['type']} | {entry['count']} | {entry['generated']} | {entry['valid']} | {entry['invalid']} | {entry['gen_speed']:.2f} | {entry['gen_time']:.2f} | {entry['val_time']:.2f} | {entry['total_time']:.2f} |
"""
        markdown += row

    # Statistics section
    markdown += "\n## Comparison\n\n"

    # Group by type
    original = [r for r in results if r['type'] == 'Original']
    optimized = [r for r in results if r['type'] == 'Optimized']

    if original and optimized:
        avg_gen_orig = sum(r['gen_speed'] for r in original) / len(original)
        avg_gen_opt = sum(r['gen_speed'] for r in optimized) / len(optimized)
        avg_val_orig = sum(r['val_time'] for r in original) / len(original)
        avg_val_opt = sum(r['val_time'] for r in optimized) / len(optimized)

        markdown += f"""
### Generation Performance
- **Original templates avg speed**: {avg_gen_orig:.2f} files/sec
- **Optimized templates avg speed**: {avg_gen_opt:.2f} files/sec
- **Speedup**: {(avg_gen_opt/avg_gen_orig):.2f}x

### Validation Performance
- **Original templates avg time**: {avg_val_orig:.2f} seconds
- **Optimized templates avg time**: {avg_val_opt:.2f} seconds
- **Reduction**: {((avg_val_orig - avg_val_opt)/avg_val_orig * 100):.1f}%

### Validity Rates
- **Original templates avg valid rate**: {sum(r['valid_rate'] for r in original) / len(original):.2f}%
- **Optimized templates avg valid rate**: {sum(r['valid_rate'] for r in optimized) / len(optimized):.2f}%
"""

    return markdown


def main():
    print("\n" + "="*60)
    print("FUZZER BENCHMARKING SCRIPT")
    print("="*60)
    print(f"Original templates: {ORIGINAL_TEMPLATES}")
    print(f"Optimized templates: {OPTIMIZED_TEMPLATES}")
    print(f"Files per template: {FUZZ_COUNT}")

    results = []

    # Process original templates
    print(f"\n{'='*60}")
    print("PHASE 1: Building and Testing Original Templates")
    print(f"{'='*60}")

    for template in ORIGINAL_TEMPLATES:
        if not build_fuzzer(template):
            print(f"Skipping {template} due to build failure")
            continue

        timestamp = run_fuzz_manager(template)
        if not timestamp:
            print(f"Skipping {template} due to fuzz failure")
            continue

        report = read_report(template, timestamp)
        if report:
            gen_data = report['generation']
            val_data = report['validation']

            results.append({
                'template': template,
                'type': 'Original',
                'count': gen_data['target_count'],
                'generated': gen_data['generated_count'],
                'valid': val_data['valid_count'],
                'invalid': val_data['invalid_count'],
                'gen_speed': gen_data['generation_speed_files_per_second'],
                'gen_time': gen_data['generation_time_seconds'],
                'val_time': val_data['validation_time_seconds'],
                'total_time': report['totals']['total_time_seconds'],
                'valid_rate': val_data['valid_rate_percent'],
            })
            print(f"✅ {template}: {gen_data['generation_speed_files_per_second']:.0f} f/s, {val_data['valid_count']}/{gen_data['generated_count']} valid")

    # Process optimized templates
    print(f"\n{'='*60}")
    print("PHASE 2: Building and Testing Optimized Templates")
    print(f"{'='*60}")

    for template in OPTIMIZED_TEMPLATES:
        if not build_fuzzer(template):
            print(f"Skipping {template} due to build failure")
            continue

        timestamp = run_fuzz_manager(template)
        if not timestamp:
            print(f"Skipping {template} due to fuzz failure")
            continue

        report = read_report(template, timestamp)
        if report:
            gen_data = report['generation']
            val_data = report['validation']

            results.append({
                'template': template,
                'type': 'Optimized',
                'count': gen_data['target_count'],
                'generated': gen_data['generated_count'],
                'valid': val_data['valid_count'],
                'invalid': val_data['invalid_count'],
                'gen_speed': gen_data['generation_speed_files_per_second'],
                'gen_time': gen_data['generation_time_seconds'],
                'val_time': val_data['validation_time_seconds'],
                'total_time': report['totals']['total_time_seconds'],
                'valid_rate': val_data['valid_rate_percent'],
            })
            print(f"✅ {template}: {gen_data['generation_speed_files_per_second']:.0f} f/s, {val_data['valid_count']}/{gen_data['generated_count']} valid")

    # Generate reports
    if results:
        print(f"\n{'='*60}")
        print("Generating Reports")
        print(f"{'='*60}")

        # Markdown report
        markdown_report = format_report_for_markdown(results)
        report_file = Path("BENCHMARK_REPORT.md")
        with open(report_file, 'w') as f:
            f.write(markdown_report)
        print(f"✅ Markdown report written to: {report_file}")

        # Word document report
        if create_word_report(results, "BENCHMARK_REPORT.docx"):
            print(f"✅ Word report written to: BENCHMARK_REPORT.docx")
        else:
            print(f"⚠️  Word report generation skipped (python-docx not available)")

        print(f"\n📊 Total results collected: {len(results)} templates")
    else:
        print("❌ No results to report")
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())
